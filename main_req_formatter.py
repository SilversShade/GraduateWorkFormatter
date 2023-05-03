from docx import Document
from docx.oxml import OxmlElement, ns
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor
from docx.shared import Cm


class MainRequirementsFormatter:

    @staticmethod
    def number_pages(doc, run):
        for p in doc.sections[0].footer.paragraphs:
            for r in p.runs:
                r.clear()

        doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.sections[0].different_first_page_header_footer = True

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    @staticmethod
    def change_font(font_name: str, paragraph):
        style = paragraph.style
        font = style.font
        font.name = font_name

    @staticmethod
    def change_font_size(font_size: int, paragraph):
        style = paragraph.style
        font = style.font
        font.size = Pt(font_size)

    @staticmethod
    def change_alignment(alignment: WD_PARAGRAPH_ALIGNMENT, paragraph):
        paragraph.alignment = alignment

    @staticmethod
    def change_font_color(color: RGBColor, paragraph):
        for run in paragraph.runs:
            run.font.color.rgb = color

    @staticmethod
    def change_line_spacing(spacing: WD_LINE_SPACING, paragraph):
        paragraph.paragraph_format.line_spacing_rule = spacing

    @staticmethod
    def change_margins(left_cm, right_cm, top_cm, bottom_cm, doc: Document):
        for section in doc.sections:
            section.left_margin = Cm(left_cm)
            section.right_margin = Cm(right_cm)
            section.top_margin = Cm(top_cm)
            section.bottom_margin = Cm(bottom_cm)

    @staticmethod
    def change_left_paragraph_indentation(indentation_cm, paragraph):
        paragraph.paragraph_format.left_indent = Cm(indentation_cm)

    @staticmethod
    def format_document(doc: Document):
        for paragraph in doc.paragraphs:
            MainRequirementsFormatter.change_font('Times New Roman', paragraph)
            MainRequirementsFormatter.change_font_size(14, paragraph)
            MainRequirementsFormatter.change_alignment(WD_PARAGRAPH_ALIGNMENT.CENTER, paragraph)
            MainRequirementsFormatter.change_font_color(RGBColor(0, 0, 0), paragraph)
            MainRequirementsFormatter.change_line_spacing(WD_LINE_SPACING.ONE_POINT_FIVE, paragraph)
            MainRequirementsFormatter.change_left_paragraph_indentation(1.25, paragraph)
        MainRequirementsFormatter.number_pages(doc, doc.sections[0].footer.paragraphs[0].add_run())
        MainRequirementsFormatter.change_margins(0.3, 0.15, 0.2, 0.2, doc)
